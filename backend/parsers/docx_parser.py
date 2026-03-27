"""
DOCX Parser — two-layer approach for robust text extraction.

Layer 1: python-docx (structured paragraph + table extraction)
Layer 2: mammoth (fallback for complex formatting, embedded objects)
"""

import io
import logging
from zipfile import BadZipFile

import mammoth
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from parsers.text_cleaner import clean_text

logger = logging.getLogger(__name__)


def parse_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using a two-layer parsing strategy.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        Cleaned text content of the DOCX.

    Raises:
        ValueError: If the DOCX cannot be parsed or is empty.
    """
    text = ""

    # ── Layer 1: python-docx (primary) ─────────────────────────
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract table content (resumes often use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        text = "\n".join(paragraphs).strip()
    except (BadZipFile, PackageNotFoundError):
        logger.warning("python-docx rejected invalid DOCX package")
    except Exception as e:
        logger.warning("python-docx failed: %s — falling back to mammoth", e)

    # ── Layer 2: mammoth fallback ──────────────────────────────
    if len(text) < 100:
        try:
            result = mammoth.extract_raw_text(io.BytesIO(file_bytes))
            text = result.value
        except BadZipFile as e:
            raise ValueError(
                "Invalid or corrupted DOCX file. Please upload a valid .docx exported from Word or Google Docs."
            ) from e
        except Exception as e:
            raise ValueError(
                "Could not parse DOCX. The file may be invalid, corrupted, or unsupported."
            ) from e

    # ── Edge case: empty document ──────────────────────────────
    if len(text.strip()) < 50:
        raise ValueError(
            "DOCX appears to be empty or contains only images."
        )

    return clean_text(text)
